"""
주간 해외매크로 보고서 생성기

사용법:
  python generate_report.py                     # data/ 폴더에서 최신 파일 자동 선택
  python generate_report.py 20260519            # 날짜 직접 지정
  python generate_report.py data/data_20260512.json  # 파일 경로 직접 지정

출력: archive/YYYY/Weekly foreign-macro_YYYYMMDD.docx
"""

import json
import sys
import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

# ── 색상 팔레트 ────────────────────────────────────────────
FONT          = "맑은 고딕"
COLOR_TITLE   = RGBColor(0x1F, 0x49, 0x7D)
COLOR_HEAD1   = RGBColor(0x2E, 0x74, 0xB5)
COLOR_HEAD2   = RGBColor(0x2E, 0x74, 0xB5)
COLOR_TH_BG   = "2E74B5"
COLOR_GRAY    = "D9D9D9"
COLOR_RED     = RGBColor(0xC0, 0x00, 0x00)
COLOR_GREEN   = RGBColor(0x37, 0x86, 0x20)


# ── 공통 헬퍼 ─────────────────────────────────────────────

def _east_asia_font(run):
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), FONT)
    rPr.insert(0, rFonts)


def set_font(run, size=10, bold=False, color=None, italic=False):
    run.font.name = FONT
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic
    if color:
        run.font.color.rgb = color
    _east_asia_font(run)


def add_heading(doc, text, level=1):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(10 if level == 1 else 6)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(text)
    set_font(run, 14 if level == 1 else 11, bold=True, color=COLOR_HEAD1)
    if level == 1:
        pPr  = p._p.get_or_add_pPr()
        pBdr = OxmlElement("w:pBdr")
        bot  = OxmlElement("w:bottom")
        bot.set(qn("w:val"),   "single")
        bot.set(qn("w:sz"),    "6")
        bot.set(qn("w:space"), "1")
        bot.set(qn("w:color"), "2E74B5")
        pBdr.append(bot)
        pPr.append(pBdr)
    return p


def add_subheading(doc, text):
    return add_heading(doc, text, level=2)


def add_bullet(doc, text, bold_prefix=None):
    p = doc.add_paragraph(style="List Bullet")
    p.paragraph_format.space_after  = Pt(1)
    p.paragraph_format.left_indent  = Cm(0.5)
    if bold_prefix:
        r1 = p.add_run(bold_prefix)
        set_font(r1, 10, bold=True)
        r2 = p.add_run(text)
        set_font(r2, 10)
    else:
        set_font(p.add_run(text), 10)
    return p


def shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd  = OxmlElement("w:shd")
    shd.set(qn("w:val"),   "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"),  hex_color)
    tcPr.append(shd)


def _value_color(text):
    """양수면 초록, 음수면 빨강, 그 외 None."""
    if text.startswith("+") or any(k in text for k in ("상회", "상승")):
        return COLOR_GREEN
    if text.startswith("-") or any(k in text for k in ("하회", "하락")):
        return COLOR_RED
    return None


def add_table(doc, headers, rows, col_widths=None):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style     = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    # 헤더 행
    for i, h in enumerate(headers):
        cell = table.rows[0].cells[i]
        shade_cell(cell, COLOR_TH_BG)
        p   = cell.paragraphs[0]
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_font(p.add_run(h), 9, bold=True, color=RGBColor(0xFF, 0xFF, 0xFF))

    # 데이터 행
    for r_idx, row_data in enumerate(rows):
        row = table.rows[r_idx + 1]
        if r_idx % 2 == 1:
            for cell in row.cells:
                shade_cell(cell, COLOR_GRAY)
        for c_idx, val in enumerate(row_data):
            cell = row.cells[c_idx]
            p    = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            text = str(val)
            set_font(p.add_run(text), 9, color=_value_color(text))

    # 열 너비
    if col_widths:
        for i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[i].width = Cm(w)

    doc.add_paragraph()
    return table


def add_notice(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.space_after  = Pt(3)
    p.paragraph_format.left_indent  = Cm(0.3)
    set_font(p.add_run("※ " + text), 9, italic=True,
             color=RGBColor(0x60, 0x60, 0x60))
    return p


# ── 섹션별 빌더 ───────────────────────────────────────────

def build_title(doc, meta):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(20)
    set_font(p.add_run("주간 해외매크로 보고"), 20, bold=True, color=COLOR_TITLE)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    date_str = datetime.date.fromisoformat(meta["date"]).strftime("%Y.%m.%d")
    set_font(sub.add_run(f"보고 기간: {meta['period']}   |   작성일: {date_str}"),
             10, color=RGBColor(0x60, 0x60, 0x60))
    doc.add_paragraph()


def build_summary(doc, data):
    add_heading(doc, "1. 주간 핵심 요약")
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(0.5)
    p.paragraph_format.right_indent = Cm(0.5)
    p.paragraph_format.space_after  = Pt(6)
    set_font(p.add_run(data["summary"]), 10.5, bold=True, color=COLOR_TITLE)


def build_markets(doc, m):
    add_heading(doc, "2. 글로벌 금융시장 동향")

    # 2-1 증시
    add_subheading(doc, "2-1. 글로벌 증시")
    add_table(doc,
        headers=["지수", "종가", "주간 등락", "YTD", "비고"],
        rows=[[s["index"], s["close"], s["weekly"], s["ytd"], s["note"]]
              for s in m["stocks"]],
        col_widths=[3.5, 2.5, 2.5, 2.0, 5.0])
    add_notice(doc, m["stocks_notice"])

    # 2-2 채권금리
    add_subheading(doc, "2-2. 미국 국채금리")
    add_table(doc,
        headers=["만기", "금주", "전주", "전주대비", "비고"],
        rows=[[b["maturity"], b["current"], b["prev_week"], b["change"], b["note"]]
              for b in m["bonds"]],
        col_widths=[3.5, 2.8, 2.8, 2.8, 5.0])
    add_notice(doc, m["bonds_notice"])

    # 2-3 원자재
    add_subheading(doc, "2-3. 주요 원자재")
    add_table(doc,
        headers=["품목", "금주 종가", "주간 등락", "전월비", "전년비", "비고"],
        rows=[[c["item"], c["close"], c["weekly"], c["mom"], c["yoy"], c["note"]]
              for c in m["commodities"]],
        col_widths=[3.8, 2.5, 2.5, 2.2, 2.2, 4.0])
    add_notice(doc, m["commodities_notice"])

    # 2-4 외환
    add_subheading(doc, "2-4. 외환(FX)")
    add_table(doc,
        headers=["통화쌍", "금주", "전월비", "전년비", "비고"],
        rows=[[f["pair"], f["current"], f["mom"], f["yoy"], f["note"]]
              for f in m["fx"]],
        col_widths=[3.5, 3.0, 2.5, 2.5, 5.5])

    # 2-5 경제지표
    add_subheading(doc, "2-5. 주간 주요 경제지표 발표")
    add_table(doc,
        headers=["지표명", "발표일", "실제치", "예상치", "전월치", "평가"],
        rows=[[i["name"], i["date"], i["actual"], i["forecast"], i["prev"], i["eval"]]
              for i in m["indicators"]],
        col_widths=[4.5, 2.0, 2.2, 2.2, 2.2, 4.0])
    add_notice(doc, m["indicators_notice"])

    # 2-6 연준
    add_subheading(doc, "2-6. 연준(Fed) 동향")
    fed = m["fed"]
    add_bullet(doc, fed["rate"],   bold_prefix="현행 기준금리: ")
    add_bullet(doc, fed["next"],   bold_prefix="다음 FOMC: ")
    add_bullet(doc, fed["gov"],    bold_prefix="지배구조: ")
    add_bullet(doc, fed["market"], bold_prefix="시장 기대: ")
    doc.add_paragraph()


def build_news(doc, news_list):
    add_heading(doc, "3. 핵심 이슈 뉴스")
    marks = ["①", "②", "③"]
    for idx, n in enumerate(news_list):
        mark = marks[idx] if idx < len(marks) else f"{idx+1}"
        add_subheading(doc, f"이슈 {mark} {n['title']}")
        add_bullet(doc, n["background"], bold_prefix="[사실] 배경: ")
        add_bullet(doc, n["progress"],   bold_prefix="[사실] 경과: ")
        add_bullet(doc, n["impact"],     bold_prefix="[사실] 시장 영향: ")
        add_bullet(doc, n["opinion"],    bold_prefix="[의견] ")
        doc.add_paragraph()


def build_next_week(doc, nw):
    add_heading(doc, f"4. 다음 주 경제지표 및 주요 일정 ({nw['title']})")
    add_table(doc,
        headers=["날짜", "시간(ET)", "국가", "지표 / 이벤트", "예상치", "전월치", "중요도"],
        rows=[[c["date"], c["time"], c["country"], c["event"],
               c["forecast"], c["prev"], c["importance"]]
              for c in nw["calendar"]],
        col_widths=[2.4, 2.0, 1.5, 5.0, 2.0, 2.0, 2.2])
    add_notice(doc, nw["calendar_notice"])

    p = doc.add_paragraph()
    set_font(p.add_run("주요 인사 발언 예정"), 10, bold=True)
    for s in nw["speakers"]:
        add_bullet(doc, s)
    doc.add_paragraph()


def build_outlook(doc, o):
    add_heading(doc, "5. 시사점 및 전망")

    add_subheading(doc, "5-1. 경기·인플레이션")
    add_bullet(doc, o["economy"]["fact"],    bold_prefix="[사실] ")
    add_bullet(doc, o["economy"]["opinion"], bold_prefix="[의견] ")

    add_subheading(doc, "5-2. 통화정책")
    add_bullet(doc, o["monetary"]["fact"],    bold_prefix="[사실] ")
    add_bullet(doc, o["monetary"]["opinion"], bold_prefix="[의견] ")

    add_subheading(doc, "5-3. 지정학·무역")
    add_bullet(doc, o["geopolitical"]["fact"],    bold_prefix="[사실] ")
    add_bullet(doc, o["geopolitical"]["opinion"], bold_prefix="[의견] ")

    add_subheading(doc, "5-4. 다음 주 체크포인트")
    for cp in o["checkpoints"]:
        add_bullet(doc, cp)
    doc.add_paragraph()


def build_disclaimer(doc):
    hr = doc.add_paragraph()
    hr.paragraph_format.space_before = Pt(12)
    set_font(hr.add_run("─" * 60), 9, color=RGBColor(0xAA, 0xAA, 0xAA))

    disc = doc.add_paragraph()
    set_font(disc.add_run(
        "본 보고서는 공개된 자료를 기반으로 정보 제공 목적으로 작성되었습니다. "
        "투자 권유 또는 의사결정의 근거로 활용 시 추가적인 검토가 필요합니다. "
        "사실(Fact)과 의견(Opinion)은 각 항목에 명시하였습니다."
    ), 8.5, italic=True, color=RGBColor(0x80, 0x80, 0x80))


# ── 메인 ──────────────────────────────────────────────────

def resolve_data_file(arg=None):
    """CLI 인수에서 데이터 파일 경로를 결정."""
    base = Path(__file__).parent / "data"
    if arg is None:
        files = sorted(base.glob("data_*.json"))
        if not files:
            raise FileNotFoundError("data/ 폴더에 data_YYYYMMDD.json 파일이 없습니다.")
        return files[-1]
    if Path(arg).exists():
        return Path(arg)
    # 날짜 문자열(YYYYMMDD)로 전달한 경우
    candidate = base / f"data_{arg}.json"
    if candidate.exists():
        return candidate
    raise FileNotFoundError(f"데이터 파일을 찾을 수 없습니다: {arg}")


def build_report(data_path):
    with open(data_path, encoding="utf-8") as f:
        data = json.load(f)

    doc = Document()
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    build_title(doc,      data["meta"])
    build_summary(doc,    data)
    build_markets(doc,    data["markets"])
    build_news(doc,       data["news"])
    build_next_week(doc,  data["next_week"])
    build_outlook(doc,    data["outlook"])
    build_disclaimer(doc)

    # 출력 경로: archive/YYYY/Weekly foreign-macro_YYYYMMDD.docx
    report_date = data["meta"]["date"].replace("-", "")
    year        = data["meta"]["date"][:4]
    out_dir     = Path(__file__).parent / "archive" / year
    out_dir.mkdir(parents=True, exist_ok=True)
    out_path    = out_dir / f"Weekly foreign-macro_{report_date}.docx"
    doc.save(out_path)
    return out_path


if __name__ == "__main__":
    arg       = sys.argv[1] if len(sys.argv) > 1 else None
    data_path = resolve_data_file(arg)
    out       = build_report(data_path)
    sys.stdout.buffer.write(f"[OK] {out}\n".encode("utf-8"))
